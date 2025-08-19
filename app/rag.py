# app/rag.py
from __future__ import annotations
import os, re, uuid, subprocess, shlex, tempfile
from typing import List, Dict, Any, Tuple

import chromadb
from fastembed import TextEmbedding
from pypdf import PdfReader

import pypdfium2 as pdfium
import pytesseract
from PIL import Image  # noqa: F401

from docx import Document as DocxDocument        # DOCX
from openpyxl import load_workbook               # XLSX
from shutil import which

from .config import settings

# ---------------------------
# Embeddings (FastEmbed)
# ---------------------------

class FastEmbedder:
    """
    Chroma requires __call__(self, input: List[str]) -> List[List[float]]
    Parameter name must be 'input'. Return plain Python lists (not NumPy).
    """
    def __init__(self, model_name: str | None = None):
        self.model_name = model_name or "BAAI/bge-small-en-v1.5"
        self.model = TextEmbedding(self.model_name)

    def __call__(self, input):
        batch = list(input)  # Chroma may pass a tuple
        return [vec.tolist() for vec in self.model.embed(batch)]

embedder = FastEmbedder(getattr(settings, "EMBEDDING_MODEL", None))

# ---------------------------
# Vector store (Chroma)
# ---------------------------

def get_chroma():
    client = chromadb.PersistentClient(path=settings.CHROMA_DIR)
    collection = client.get_or_create_collection(
        name='ezzogenics_kb',
        metadata={'hnsw:space': 'cosine'},
        embedding_function=embedder
    )
    return collection

# ---------------------------
# Text utils
# ---------------------------

def clean_text(t: str) -> str:
    t = re.sub(r'\s+', ' ', t)
    return t.strip()

def split_text(text: str, max_len: int = 500, overlap: int = 120) -> List[str]:
    """Paragraph/sentence chunker with overlap."""
    paras = re.split(r'(?:\n\s*\n)+', text)
    chunks: List[str] = []
    buf: List[str] = []
    size = 0
    for p in paras:
        p = p.strip()
        if not p:
            continue
        sents = re.split(r'(?<=[.!?])\s+', p)
        for s in sents:
            if size + len(s) > max_len and buf:
                chunks.append(' '.join(buf))
                joined = ' '.join(buf)
                tail = joined[-overlap:] if overlap > 0 else ''
                buf = [tail, s] if tail else [s]
                size = len(tail) + len(s)
            else:
                buf.append(s)
                size += len(s)
    if buf:
        chunks.append(' '.join(buf))
    return [clean_text(c) for c in chunks if c.strip()]

# ---------------------------
# PDF parsing (+ OCR fallback)
# ---------------------------

def ocr_pdf_page(pdf_path: str, page_index: int) -> str:
    pdf = pdfium.PdfDocument(pdf_path)
    page = pdf.get_page(page_index)
    pil_img = page.render(scale=2).to_pil()
    page.close()
    text = pytesseract.image_to_string(pil_img, lang="eng")
    return clean_text(text or "")

def parse_pdf(path: str) -> List[Tuple[str, int]]:
    out: List[Tuple[str, int]] = []
    reader = PdfReader(path)
    for i, _ in enumerate(reader.pages):
        page_num = i + 1
        try:
            raw = reader.pages[i].extract_text() or ""
        except Exception:
            raw = ""
        text = clean_text(raw)
        if len(text) < 30:
            try:
                text = ocr_pdf_page(path, i)
            except Exception as e:
                print(f"[WARN] OCR failed for {os.path.basename(path)} p.{page_num}: {e}")
                text = ""
        if text:
            out.append((text, page_num))
    return out

# ---------------------------
# Office parsing (DOCX / XLSX)
# ---------------------------

def parse_docx(path: str) -> List[Tuple[str, int]]:
    """Return 1 'page' per ~1500 chars so it behaves like paginated text."""
    doc = DocxDocument(path)
    lines: List[str] = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            lines.append(p.text.strip())
    # tables (simple flatten)
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [clean_text(c.text or "") for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    full = "\n".join(lines)
    if not full.strip():
        return []
    chunks = []
    step = 1500
    for i in range(0, len(full), step):
        part = full[i:i+step]
        chunks.append((clean_text(part), i // step + 1))
    return chunks

def parse_xlsx(path: str) -> List[Tuple[str, int]]:
    """Each sheet becomes one 'page' of concatenated rows."""
    wb = load_workbook(path, read_only=True, data_only=True)
    out: List[Tuple[str, int]] = []
    for si, sheet_name in enumerate(wb.sheetnames, start=1):
        sh = wb[sheet_name]
        rows: List[str] = []
        for row in sh.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                rows.append(" | ".join(cells))
        text = clean_text(f"Sheet: {sheet_name}\n" + "\n".join(rows))
        if text:
            out.append((text, si))
    return out

# ---------------------------
# .DOC/.DOCM conversion via LibreOffice (temp only; no artifacts)
# ---------------------------

def _find_soffice() -> str | None:
    # 1) Env var override
    env = os.getenv("SOFFICE_PATH")
    if env and os.path.exists(env):
        return env
    # 2) PATH
    p = which("soffice")
    if p:
        return p
    # 3) macOS default for LibreOffice
    mac = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
    if os.path.exists(mac):
        return mac
    return None

def convert_doc_to_pages(path: str) -> List[Tuple[str, int]]:
    """
    Convert legacy .doc/.docm to PDF in a temp dir, parse pages, return text.
    No *.converted.pdf files are written to your data folder.
    """
    soffice = _find_soffice()
    if not soffice:
        print("[WARN] LibreOffice 'soffice' not found. Set SOFFICE_PATH or install LibreOffice.")
        return []

    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            cmd = f'{shlex.quote(soffice)} --headless --convert-to pdf --outdir {shlex.quote(tmpdir)} {shlex.quote(path)}'
            subprocess.run(cmd, check=True, shell=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            base = os.path.splitext(os.path.basename(path))[0]
            pdf_out = os.path.join(tmpdir, base + ".pdf")
            if os.path.exists(pdf_out):
                return parse_pdf(pdf_out)
            else:
                print(f"[WARN] LibreOffice produced no PDF for {os.path.basename(path)}")
                return []
    except Exception as e:
        print(f"[WARN] LibreOffice conversion failed for {os.path.basename(path)}: {e}")
        return []

# ---------------------------
# Ingestion
# ---------------------------

def ingest_folder(data_dir: str = './data') -> Dict[str, Any]:
    col = get_chroma()
    added = 0
    if not os.path.isdir(data_dir):
        return {'chunks_added': 0, 'note': f'No data dir: {data_dir}'}

    for name in os.listdir(data_dir):
        # Skip previously-generated artifacts if any still linger
        if name.lower().endswith('.converted.pdf'):
            continue

        path = os.path.join(data_dir, name)
        low = name.lower()

        try:
            if low.endswith('.pdf'):
                pages = parse_pdf(path)
            elif low.endswith('.docx'):
                pages = parse_docx(path)
            elif low.endswith('.xlsx'):
                pages = parse_xlsx(path)
            elif low.endswith('.doc') or low.endswith('.docm'):
                pages = convert_doc_to_pages(path)
                if not pages:
                    print(f"[WARN] Skipping {name}: cannot convert to PDF and no direct parser.")
                    continue
            else:
                print(f"[INFO] Skipping unsupported file type: {name}")
                continue

        except Exception as e:
            print(f'[WARN] Could not parse {name}: {e}')
            continue

        for page_text, page_num in pages:
            for chunk in split_text(page_text):
                doc_id = str(uuid.uuid4())
                payload = f"{name} | page {page_num}\n{chunk}"
                col.add(
                    ids=[doc_id],
                    documents=[payload],
                    metadatas=[{'source': name, 'page': page_num}]
                )
                added += 1

    return {'chunks_added': added}

# ---------------------------
# Retrieval
# ---------------------------

def retrieve(query: str, top_k: int | None = None):
    col = get_chroma()
    k = top_k or settings.TOP_K
    res = col.query(
        query_texts=[query],
        n_results=k,
        include=['documents', 'metadatas', 'distances']
    )
    docs = res.get('documents', [[]])[0]
    metas = res.get('metadatas', [[]])[0]
    dists = res.get('distances', [[]])[0]
    items = []
    for d, m, dist in zip(docs, metas, dists):
        items.append({
            'text': d,
            'source': m.get('source'),
            'page': m.get('page'),
            'score': 1.0 - float(dist) if dist is not None else None
        })
    return items

# ---------------------------
# LLM calls
# ---------------------------

SYS_PROMPT = (
    """You are a helpful analyst for Ezzogenics. Answer ONLY using the provided context.

If the answer is not present in the context, say you don't have enough information and suggest where to look.
Always include a short bullet list of key facts you used.
Include no speculative claims. Keep answers concise and specific.
"""
)

def build_prompt(query: str, contexts: List[Dict[str, Any]]) -> List[Dict[str, str]]:
    if not contexts:
        user = f"No relevant context found.\n\nQuestion: {query}\n\nAnswer:"
    else:
        context_text = "\n\n".join([f"[Source: {c['source']} p.{c['page']}]\n{c['text']}" for c in contexts])
        user = f"Context:\n{context_text}\n\nQuestion: {query}\n\nAnswer:"
    return [
        {"role": "system", "content": SYS_PROMPT},
        {"role": "user", "content": user},
    ]

def call_llm(messages: List[Dict[str, str]]) -> str:
    provider = (settings.LLM_PROVIDER or "").lower()
    if provider == 'openai':
        from openai import OpenAI
        client = OpenAI(api_key=settings.OPENAI_API_KEY)
        resp = client.chat.completions.create(
            model=settings.OPENAI_MODEL,
            messages=messages,
            temperature=settings.TEMPERATURE,
            max_tokens=settings.MAX_TOKENS,
        )
        return resp.choices[0].message.content.strip()
    else:
        from groq import Groq
        client = Groq(api_key=settings.GROQ_API_KEY)
        resp = client.chat.completions.create(
            model=settings.GROQ_MODEL,
            messages=messages,
            temperature=settings.TEMPERATURE,
            max_tokens=settings.MAX_TOKENS,
        )
        return resp.choices[0].message.content.strip()

def answer_query(query: str) -> Dict[str, Any]:
    ctx = retrieve(query)
    messages = build_prompt(query, ctx)
    answer = call_llm(messages)
    cites = [{'source': c.get('source'), 'page': c.get('page')} for c in ctx] if ctx else []
    return {'answer': answer, 'citations': cites}
